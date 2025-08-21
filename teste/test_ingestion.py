"""Teste a extração do '.docx' e o chunking."""
import os 
import docx
from modules.ingestion import extract_chapters_from_word, get_text_chunks
import config

def test_extract_chapters_from_word(tmpdir):
    """Teste se a função consegue extrair capítulos de um '.docx' mock."""
    # Cria um arquivo '.docx' falso para o teste
    doc = docx.Document()
    doc.add_paragraph("Introdução geral do livro.")
    doc.add_paragraph("Unidade 1 - Capítulo 1")
    doc.add_paragraph("Texto do primeiro capítulo.")
    doc.add_paragraph("Unidade 1 - Capítulo 2")
    doc.add_paragraph("Texto do segundo capítulo.")

    test_docx_path = os.path.join(tmpdir, "test_book.docx")
    doc.save(test_docx_path)

    chapters = extract_chapters_from_word(test_docx_path)

    assert isinstance(chapters, dict)
    assert len(chapters) == 2
    assert (1, 1) in chapters
    assert "Texto do primeiro capítulo." in chapters[(1, 1)]
    assert "Texto do segundo capítulo." in chapters[(1, 2)]

def test_get_text_chunks(sample_chapter_text):
    """Testa se o texto é dividido em uma lista de strings."""
    chunks = get_text_chunks(sample_chapter_text)
    assert isinstance(chunks, list)
    assert len(chunks) > 0
    assert all(isinstance(chunk, str) for chunk in chunks)

def test_show_content_of_first_chapter():
    """
    Teste de validação para extrair e exibir o INÍCIO e o FIM do conteúdo
    do primeiro capítulo (U1, C1) para inspeção manual.
    """
    print("\n--- INICIANDO TESTE DE VALIDAÇÃO: EXTRAÇÃO DO CAPÍTULO 1 ---")
    
    # 1. Tenta extrair todos os capítulos do livro definido no config
    chapters_data = extract_chapters_from_word(config.LIVRO_ORIGINAL_PATH)
    
    # Garante que a extração retornou algo
    assert chapters_data, "A extração de capítulos falhou ou não retornou dados."

    # 2. Procura pelo primeiro capítulo (Unidade 1, Capítulo 1)
    first_chapter_data = None
    for chapter in chapters_data:
        if chapter.get("unit") == 1 and chapter.get("chapter") == 1:
            first_chapter_data = chapter
            break
            
    # Garante que o capítulo foi encontrado
    assert first_chapter_data is not None, "Unidade 1, Capítulo 1 não foi encontrado no documento."
    
    first_chapter_text = first_chapter_data.get("text", "")
    title = first_chapter_data.get("title", "Título não encontrado")
    
    # Garante que o texto do capítulo não está vazio
    assert len(first_chapter_text) > 0, "O texto do Capítulo 1 está vazio."

    # 3. Imprime o conteúdo para validação manual
    print("\n" + "="*80)
    print(f"CONTEÚDO EXTRAÍDO PARA: Unidade 1 - Capítulo 1: '{title}'")
    print("="*80)
    
    # --- INÍCIO DA MODIFICAÇÃO ---
    
    # Imprime os primeiros 1000 caracteres
    print("--- INÍCIO DO CONTEÚDO DO CAPÍTULO ---")
    print(first_chapter_text[:1000] + "...")
    
    print("\n" + "..."*20 + "\n")
    
    # Imprime os últimos 1000 caracteres
    print("--- FINAL DO CONTEÚDO DO CAPÍTULO ---")
    print("..." + first_chapter_text[-8000:])

    # --- FIM DA MODIFICAÇÃO ---
    
    print("="*80)
    print("VALIDAÇÃO: Verifique se o texto de INÍCIO e FIM correspondem ao conteúdo esperado.")
    print("-" * 80 + "\n")